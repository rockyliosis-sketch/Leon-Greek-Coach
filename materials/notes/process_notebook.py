import fitz # PyMuPDF
import os
import re
import json
import time
import google.generativeai as genai
from PIL import Image
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 1. Config paths
PROJECT_DIR = "/Users/johnsmacbook/Documents/antigravity IDE/Projects/Leon-Greek-Coach"
BASE_DIR = "/Users/johnsmacbook/Documents/antigravity IDE/Greek book/希腊语学习笔记"
PDF_PATH = os.path.join(BASE_DIR, "notebook.pdf")
PROGRESS_PATH = os.path.join(BASE_DIR, "transcribe_progress.json")
OUTPUT_DOCX = os.path.join(BASE_DIR, "希腊语学习笔记.docx")

# 2. Configure GenAI API
load_dotenv(os.path.join(PROJECT_DIR, ".env"))
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("GEMINI_API_KEY is not configured in .env file.")
genai.configure(api_key=api_key)

# Helper to call Gemini with retries
def call_gemini_ocr(image_path, page_idx, retries=5, delay=5):
    prompt = (
        "请识别图片中手写的希腊语学习笔记，并返回一个 JSON 对象，包含以下两个字段：\n"
        "1. 'date': 字符串。如果在页面中找到了写下的学习日期，请将日期进行标准化输出。例如：如果写的是 '7/9/25' 或 '7/9'，由于第一次笔记是 2025 年 9 月 7 日，代表 2025 年 9 月 7 日，请标准化输出为 '2025-09-07'；如果是 '14/9'，请标准化输出为 '2025-09-14'；若跨年到 1 月或后面，请根据上下文推断年份（如 2026 年）并输出。如果没有找到日期，请返回 null。\n"
        "2. 'text': 字符串。请识别图片中的全部手写内容，转录为清晰易读的文本。希腊字母或单词要准确记录，如果有中文或英文解释也一并转录。保持原始笔记的段落和列表格式。如果有看不清楚的字词，请根据前后文及希腊语常识自行理解并补充完整，以确保笔记的连贯性和准确性。\n"
        "请直接输出 JSON，不要使用 markdown 标记包裹 JSON 代码块。"
    )
    
    for attempt in range(retries):
        try:
            image = Image.open(image_path)
            model = genai.GenerativeModel('gemini-flash-lite-latest')
            response = model.generate_content(
                [prompt, image],
                generation_config={"response_mime_type": "application/json"}
            )
            if response and response.text:
                # Try parsing JSON to validate
                data = json.loads(response.text)
                if 'text' in data:
                    return data
                else:
                    print(f"[Page {page_idx+1}] JSON parsed but key 'text' is missing. Retry.")
            else:
                print(f"[Page {page_idx+1}] Empty response from Gemini API.")
        except Exception as e:
            print(f"[Page {page_idx+1}] Error on attempt {attempt+1}/{retries}: {e}")
            if attempt < retries - 1:
                time.sleep(delay * (attempt + 1))
    return None

# Normalize date strings helper
def normalize_date_string(date_str, current_year=2025):
    if not date_str:
        return None, current_year
    date_str = date_str.strip()
    
    # Check YYYY-MM-DD
    match = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})$", date_str)
    if match:
        y, m, d = int(match.group(1)), int(match.group(2)), int(match.group(3))
        return f"{y}年{m}月{d}日", y

    # Check DD/MM/YY or DD/MM/YYYY
    match = re.match(r"^(\d{1,2})[/\.-](\d{1,2})[/\.-](\d{2,4})$", date_str)
    if match:
        d, m, y = int(match.group(1)), int(match.group(2)), int(match.group(3))
        if y < 100:
            y += 2000
        return f"{y}年{m}月{d}日", y

    # Check DD/MM
    match = re.match(r"^(\d{1,2})[/\.-](\d{1,2})$", date_str)
    if match:
        d, m = int(match.group(1)), int(match.group(2))
        y = 2025
        if m < 8: # If month is Jan to Jul, it's 2026
            y = 2026
        return f"{y}年{m}月{d}日", y

    # Check Chinese patterns like "9月7日" or "2025年9月7日"
    match = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", date_str)
    if match:
        y, m, d = int(match.group(1)), int(match.group(2)), int(match.group(3))
        return f"{y}年{m}月{d}日", y

    match = re.search(r"(\d{1,2})月(\d{1,2})日", date_str)
    if match:
        m, d = int(match.group(1)), int(match.group(2))
        y = 2025
        if m < 8:
            y = 2026
        return f"{y}年{m}月{d}日", y

    return date_str, current_year

# Helper to set run fonts (supporting Chinese & Greek/English)
def set_run_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

def main():
    print("Opening PDF...")
    doc = fitz.open(PDF_PATH)
    total_pages = len(doc)
    print(f"Total pages to process: {total_pages}")
    
    # Load progress
    if os.path.exists(PROGRESS_PATH):
        with open(PROGRESS_PATH, 'r', encoding='utf-8') as f:
            progress = json.load(f)
        print(f"Loaded existing progress. {len(progress.get('processed_pages', {}))} pages already processed.")
    else:
        progress = {
            "processed_pages": {},
            "metadata": {"total_pages": total_pages}
        }
        
    # Process pages
    for idx in range(total_pages):
        page_key = str(idx)
        if page_key in progress["processed_pages"]:
            continue
            
        print(f"[{idx+1}/{total_pages}] Processing page...")
        
        # 1. Render page to image
        page = doc[idx]
        mat = fitz.Matrix(0.25, 0.25)
        pix = page.get_pixmap(matrix=mat)
        temp_img_path = os.path.join(BASE_DIR, f"temp_page_{idx}.png")
        pix.save(temp_img_path)
        
        # 2. Call Gemini
        result = call_gemini_ocr(temp_img_path, idx)
        
        # Clean up image immediately
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)
            
        if result:
            progress["processed_pages"][page_key] = result
            print(f"[{idx+1}/{total_pages}] Successfully transcribed. Date: {result.get('date')}")
            
            # Save progress
            with open(PROGRESS_PATH, 'w', encoding='utf-8') as f:
                json.dump(progress, f, ensure_ascii=False, indent=2)
        else:
            print(f"[{idx+1}/{total_pages}] FAILED to transcribe page. Stopping batch to avoid missing content.")
            return

        # Avoid hitting rate limits
        time.sleep(1.5)

    print("\nAll pages successfully transcribed! Compiling Word Document...")
    
    # Compile DOCX
    docx_doc = Document()
    
    # Set document margins (1 inch)
    for section in docx_doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Group content by date
    # Start date defaults to "2025年9月7日" (First note date)
    current_date = "2025年9月7日"
    current_year = 2025
    date_groups = {} # date_str -> list of (page_num, text)
    
    for idx in range(total_pages):
        page_data = progress["processed_pages"][str(idx)]
        raw_date = page_data.get("date")
        
        if raw_date:
            normalized, year = normalize_date_string(raw_date, current_year)
            if normalized:
                current_date = normalized
                current_year = year
                
        if current_date not in date_groups:
            date_groups[current_date] = []
            
        date_groups[current_date].append((idx + 1, page_data["text"]))

    # Write date groups to docx
    for date_str, pages in date_groups.items():
        # Add Date Heading
        h = docx_doc.add_heading(level=1)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
        h.paragraph_format.keep_with_next = True
        
        run = h.add_run(date_str)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(43, 87, 154) # Classic Slate Blue
        set_run_font(run)
        
        for page_num, text in pages:
            # Page separator caption
            p_cap = docx_doc.add_paragraph()
            p_cap.paragraph_format.space_before = Pt(12)
            p_cap.paragraph_format.space_after = Pt(4)
            p_cap.paragraph_format.keep_with_next = True
            
            run_cap = p_cap.add_run(f"—— 原笔记第 {page_num} 页 ——")
            run_cap.italic = True
            run_cap.font.size = Pt(9.5)
            run_cap.font.color.rgb = RGBColor(128, 128, 128)
            set_run_font(run_cap)
            
            # Add text lines
            lines = text.split("\n")
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue
                
                # Check for bullet items
                if stripped.startswith(("-", "*", "•")):
                    p = docx_doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_after = Pt(4)
                    content = stripped.lstrip("-*• ").strip()
                    run_line = p.add_run(content)
                else:
                    p = docx_doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    run_line = p.add_run(line)
                    
                run_line.font.size = Pt(11)
                set_run_font(run_line)
                
    docx_doc.save(OUTPUT_DOCX)
    print(f"\nSuccessfully generated Word Document: {OUTPUT_DOCX}")

if __name__ == "__main__":
    main()
