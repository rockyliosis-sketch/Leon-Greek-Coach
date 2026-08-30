import fitz  # PyMuPDF
import os
import re
import json
import time
import sys
import google.generativeai as genai
from PIL import Image
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Config paths
PROJECT_DIR = "/Users/johnsmacbook/Documents/antigravity IDE/Projects/Leon-Greek-Coach"
BASE_DIR = "/Users/johnsmacbook/Documents/antigravity IDE/Greek book"

# Configure GenAI API
load_dotenv(os.path.join(PROJECT_DIR, ".env"))
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("GEMINI_API_KEY is not configured in the project .env file.")
genai.configure(api_key=api_key)

# Helper to call Gemini with retries and exponential backoff
def call_gemini_ocr_translation(image_path, page_idx, retries=5, delay=5):
    prompt = (
        "你是一个专业的希腊语教师和翻译。这是一张希腊语教材页面的照片。请你把这页的内容完整抓取，并整理为抓取文字信息和教学资料的结构化文本：\n"
        "1. 页面中如果含有课文、对话、单词表、语法解释或练习题，请清晰地分段转录。\n"
        "2. 请保证所有希腊语字符的绝对准确（包括重音符号，如 ά, έ, ί, ό, ύ, ώ, ή 等）。\n"
        "3. 请务必为所有的希腊语课文、对话、短语和单词配上高质量的中文翻译，以利于Leon Greek Coach学习网站的后台识别和学习者使用。\n"
        "4. 如果有一些看不清楚的字母，或者由于印刷/拍摄质量导致的问题，请结合上下文和希腊语教材的语法、词汇常识，进行合理的理解并补充完整。\n"
        "5. 结构要清晰，比如：【课文/对话】、【生词列表】、【语法解析】、【练习题】，以便制作成教学资料。\n"
        "请直接输出整理后的Markdown文本，不要返回任何JSON包装，也不要包含任何额外的开场白或解释。"
    )
    
    for attempt in range(retries):
        try:
            image = Image.open(image_path)
            model = genai.GenerativeModel('gemini-flash-lite-latest')
            response = model.generate_content([prompt, image])
            if response and response.text:
                return response.text.strip()
            else:
                print(f"[Page {page_idx+1}] Empty response from Gemini API. Attempt {attempt+1}/{retries}")
        except Exception as e:
            print(f"[Page {page_idx+1}] Error on attempt {attempt+1}/{retries}: {e}")
            if attempt < retries - 1:
                time.sleep(delay * (attempt + 1))
    return None

# Helper to set run fonts (supporting Chinese & Greek/English)
def set_run_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

# Helper to parse Markdown and write formatted content to DOCX
def add_formatted_text(paragraph, text_line, font_size=11, color_rgb=None):
    # Regex to extract bold portions (**bold**)
    parts = re.split(r'(\*\*.*?\*\*)', text_line)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.size = Pt(font_size)
        if color_rgb:
            run.font.color.rgb = color_rgb
        set_run_font(run)

def compile_docx(progress, output_docx_path):
    print(f"Compiling Word Document to {output_docx_path}...")
    doc = Document()
    
    # Set document margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    processed_pages = progress.get("processed_pages", {})
    sorted_page_keys = sorted(processed_pages.keys(), key=lambda x: int(x), reverse=True)
    
    # Document Title
    title_text = os.path.splitext(os.path.basename(output_docx_path))[0]
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(24)
    title_p.paragraph_format.alignment = 1  # Center
    title_run = title_p.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(43, 87, 154)
    set_run_font(title_run)
    
    total_pages = int(progress.get("metadata", {}).get("total_pages", len(processed_pages)))
    for page_key in sorted_page_keys:
        page_num = total_pages - int(page_key)
        page_content = processed_pages[page_key]
        
        # Add Page Heading Separator
        sep_p = doc.add_paragraph()
        sep_p.paragraph_format.space_before = Pt(18)
        sep_p.paragraph_format.space_after = Pt(12)
        sep_p.paragraph_format.keep_with_next = True
        sep_run = sep_p.add_run(f"—— 原书第 {page_num} 页 ——")
        sep_run.bold = True
        sep_run.italic = True
        sep_run.font.size = Pt(11)
        sep_run.font.color.rgb = RGBColor(128, 128, 128)
        set_run_font(sep_run)
        
        lines = page_content.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            
            # Check Markdown headers
            if stripped.startswith("### "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                content = stripped[4:]
                add_formatted_text(p, content, font_size=13, color_rgb=RGBColor(43, 87, 154))
            elif stripped.startswith("## "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                content = stripped[3:]
                add_formatted_text(p, content, font_size=14, color_rgb=RGBColor(43, 87, 154))
            elif stripped.startswith("# "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_with_next = True
                content = stripped[2:]
                add_formatted_text(p, content, font_size=16, color_rgb=RGBColor(43, 87, 154))
            # Bullet items
            elif stripped.startswith(("- ", "* ")):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(4)
                content = stripped[2:]
                add_formatted_text(p, content, font_size=11)
            # Numbered items
            elif re.match(r"^\d+\.\s+", stripped):
                match = re.match(r"^(\d+)\.\s+(.*)", stripped)
                num = match.group(1)
                content = match.group(2)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                # Tab indentation simulation
                p.paragraph_format.left_indent = Inches(0.25)
                run_num = p.add_run(f"{num}. ")
                run_num.font.size = Pt(11)
                set_run_font(run_num)
                add_formatted_text(p, content, font_size=11)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                add_formatted_text(p, line, font_size=11)
                
    doc.save(output_docx_path)
    print(f"Word Document saved successfully: {output_docx_path}")

def compile_md(progress, output_md_path):
    print(f"Compiling Markdown Document to {output_md_path}...")
    processed_pages = progress.get("processed_pages", {})
    sorted_page_keys = sorted(processed_pages.keys(), key=lambda x: int(x), reverse=True)
    
    title_text = os.path.splitext(os.path.basename(output_md_path))[0]
    md_lines = [f"# {title_text}", ""]
    
    total_pages = int(progress.get("metadata", {}).get("total_pages", len(processed_pages)))
    for page_key in sorted_page_keys:
        page_num = total_pages - int(page_key)
        page_content = processed_pages[page_key]
        
        md_lines.append(f"## Page {page_num}")
        md_lines.append(f"—— 原书第 {page_num} 页 ——")
        md_lines.append("")
        md_lines.append(page_content)
        md_lines.append("")
        
    with open(output_md_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_lines))
    print(f"Markdown Document saved successfully: {output_md_path}")

def process_pdf(pdf_name, max_pages=None):
    pdf_path = os.path.join(BASE_DIR, pdf_name)
    base_name = os.path.splitext(pdf_name)[0]
    progress_path = os.path.join(BASE_DIR, f"{base_name}_progress.json")
    output_docx = os.path.join(BASE_DIR, f"{base_name}.docx")
    output_md = os.path.join(BASE_DIR, f"{base_name}.md")
    
    if not os.path.exists(pdf_path):
        print(f"Error: PDF not found at {pdf_path}")
        return False
        
    print(f"\n==========================================")
    print(f"Processing PDF: {pdf_name}")
    print(f"==========================================")
    
    doc = fitz.open(pdf_path)
    total_pages = len(doc)
    print(f"Total pages: {total_pages}")
    
    if max_pages:
        total_pages = min(total_pages, max_pages)
        print(f"Configured to process up to {total_pages} pages.")
        
    # Load progress
    if os.path.exists(progress_path):
        with open(progress_path, 'r', encoding='utf-8') as f:
            progress = json.load(f)
        print(f"Loaded existing progress. {len(progress.get('processed_pages', {}))} pages already completed.")
    else:
        progress = {
            "processed_pages": {},
            "metadata": {"total_pages": len(doc)}
        }
        
    for idx in range(total_pages):
        page_key = str(idx)
        if page_key in progress["processed_pages"]:
            continue
            
        print(f"[{idx+1}/{total_pages}] Rendering & Transcribing page...")
        
        # 1. Render page to image (0.4 scale)
        page = doc[idx]
        mat = fitz.Matrix(0.4, 0.4)
        pix = page.get_pixmap(matrix=mat)
        temp_img_path = os.path.join(BASE_DIR, f"temp_{base_name}_page_{idx}.png")
        pix.save(temp_img_path)
        
        # 2. Call Gemini OCR & Translation
        text = call_gemini_ocr_translation(temp_img_path, idx)
        
        # Clean up image immediately
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)
            
        if text:
            progress["processed_pages"][page_key] = text
            print(f"[{idx+1}/{total_pages}] Successfully transcribed.")
            
            # Save progress after every page to prevent data loss
            with open(progress_path, 'w', encoding='utf-8') as f:
                json.dump(progress, f, ensure_ascii=False, indent=2)
        else:
            print(f"[{idx+1}/{total_pages}] FAILED to transcribe page. Stopping execution.")
            doc.close()
            return False
            
        # Wait to respect rate limits (approx 4 seconds per page total)
        time.sleep(2.0)
        
    doc.close()
    
    # Compile outputs
    compile_docx(progress, output_docx)
    compile_md(progress, output_md)
    print(f"Finished processing {pdf_name}!")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python process_textbooks.py <pdf_filename> [max_pages]")
        sys.exit(1)
        
    target_pdf = sys.argv[1]
    pages_limit = int(sys.argv[2]) if len(sys.argv) > 2 else None
    
    success = process_pdf(target_pdf, pages_limit)
    if not success:
        sys.exit(1)
